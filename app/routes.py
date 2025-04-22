from docx import Document
from flask import Blueprint, render_template, request, redirect, url_for, send_file, session
from .models import Mural, Feedback
from . import db
import random
import io

main = Blueprint('main', __name__)

def gerar_id_unico():
    while True:
        codigo = ''.join(random.choices('0123456789', k=6))
        if not Mural.query.filter_by(codigo=codigo).first():
            return codigo

@main.route('/')
def index():
    return render_template('index.html')

@main.route('/create', methods=['POST'])
def create_mural():
    titulo = request.form.get('titulo')
    codigo = gerar_id_unico()
    novo_mural = Mural(codigo=codigo, titulo=titulo)
    db.session.add(novo_mural)
    db.session.commit()
    session[f'admin_{codigo}'] = True
    return redirect(url_for('main.mural', mural_id=codigo))

@main.route('/mural/<mural_id>')
def mural(mural_id):
    mural = Mural.query.filter_by(codigo=mural_id).first()
    if not mural:
        return "Mural não encontrado", 404
    feedbacks = Feedback.query.filter_by(mural_id=mural.id).order_by(Feedback.id.desc()).all()
    return render_template('mural.html', mural_id=mural.codigo, titulo=mural.titulo, feedbacks=feedbacks)

@main.route('/send/<mural_id>', methods=['POST'])
def send_feedback(mural_id):
    mural = Mural.query.filter_by(codigo=mural_id).first()
    if not mural:
        return "Mural não encontrado", 404
    texto = request.form.get('feedback')
    novo_feedback = Feedback(texto=texto, mural_id=mural.id)
    db.session.add(novo_feedback)
    db.session.commit()
    return redirect(url_for('main.mural', mural_id=mural.codigo))

@main.route('/delete/<mural_id>', methods=['POST'])
def delete_mural(mural_id):
    if not session.get(f'admin_{mural_id}'):
        return "Acesso negado", 403
    mural = Mural.query.filter_by(codigo=mural_id).first()
    if not mural:
        return "Mural não encontrado", 404
    db.session.delete(mural)
    db.session.commit()
    session.pop(f'admin_{mural_id}', None)
    return "Mural encerrado e todos os feedbacks foram deletados com sucesso."

@main.route('/mural/<mural_id>/exportar')
def exportar_mural(mural_id):
    mural = Mural.query.filter_by(codigo=mural_id).first()
    if not mural:
        return "Mural não encontrado", 404

    feedbacks = Feedback.query.filter_by(mural_id=mural.id).order_by(Feedback.id.desc()).all()

    document = Document()
    document.add_heading(f'Mural: {mural.titulo}', 0)
    document.add_paragraph('Feedbacks recebidos:', style='Heading 1')

    if feedbacks:
        for f in feedbacks:
            document.add_paragraph(f.texto, style='List Bullet')
    else:
        document.add_paragraph("Nenhum feedback recebido.")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"mural_{mural_id}.docx"
    )